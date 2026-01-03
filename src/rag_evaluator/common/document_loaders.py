import logging
from abc import ABC, abstractmethod
from collections.abc import Callable
from dataclasses import dataclass
from pathlib import Path
from typing import Any

try:
    import pypdf

    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


@dataclass
class Document:
    """Represents a loaded document"""

    content: str
    metadata: dict[str, Any]
    source: str


class DocumentLoader(ABC):
    """Abstract base class for document loaders"""

    @abstractmethod
    def load(self, file_path: str) -> Document:
        """Load document from file path"""
        pass

    @abstractmethod
    def supported_extensions(self) -> list[str]:
        """Return list of supported file extensions"""
        pass


class TXTLoader(DocumentLoader):
    """Load plain text files"""

    def load(self, file_path: str) -> Document:
        try:
            with open(file_path, encoding="utf-8") as f:
                content = f.read()
            return Document(
                content=content, metadata={"format": "txt", "size": len(content)}, source=file_path
            )
        except Exception as e:
            raise ValueError(f"Failed to load TXT {file_path}: {e}")

    def supported_extensions(self) -> list[str]:
        return [".txt"]


class PDFLoader(DocumentLoader):
    """Load PDF files using pypdf"""

    def load(self, file_path: str) -> Document:
        if not PYPDF_AVAILABLE:
            raise ImportError("pypdf is not installed. Please install it to load PDF files.")

        try:
            with open(file_path, "rb") as f:
                reader = pypdf.PdfReader(f)

                # Check if encrypted
                if reader.is_encrypted:
                    # Try to parse without password if possible, or raise error
                    try:
                        reader.decrypt("")
                    except Exception:
                        pass
                    if reader.is_encrypted:
                        raise ValueError(f"PDF is encrypted: {file_path}")

                # Extract text from all pages
                content_parts = []
                for page_num, page in enumerate(reader.pages):
                    try:
                        text = page.extract_text()
                        if text and text.strip():  # Skip empty pages
                            content_parts.append(text)
                    except Exception as e:
                        logger.warning(f"Failed to extract page {page_num} in {file_path}: {e}")

                content = "\n\n".join(content_parts)

                # Extract metadata
                metadata = {
                    "format": "pdf",
                    "pages": len(reader.pages),
                    "size": len(content),
                }

                # Add PDF metadata if available
                if reader.metadata:
                    if reader.metadata.title:
                        metadata["title"] = str(reader.metadata.title)
                    if reader.metadata.author:
                        metadata["author"] = str(reader.metadata.author)

                return Document(content=content, metadata=metadata, source=file_path)
        except Exception as e:
            raise ValueError(f"Failed to load PDF {file_path}: {e}")

    def supported_extensions(self) -> list[str]:
        return [".pdf"]


class DOCXLoader(DocumentLoader):
    """Load DOCX files using python-docx"""

    def load(self, file_path: str) -> Document:
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed. Please install it to load DOCX files.")

        try:
            doc = DocxDocument(file_path)

            # Extract text from paragraphs
            content_parts = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    content_parts.append(text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        content_parts.append(row_text)

            content = "\n\n".join(content_parts)

            # Metadata
            metadata = {
                "format": "docx",
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "size": len(content),
            }

            # Core properties if available
            if hasattr(doc, "core_properties"):
                if doc.core_properties.title:
                    metadata["title"] = str(doc.core_properties.title)
                if doc.core_properties.author:
                    metadata["author"] = str(doc.core_properties.author)

            return Document(content=content, metadata=metadata, source=file_path)
        except Exception as e:
            raise ValueError(f"Failed to load DOCX {file_path}: {e}")

    def supported_extensions(self) -> list[str]:
        return [".docx"]


def create_loader(file_path: str) -> DocumentLoader:
    """Factory function to create appropriate loader based on file extension"""
    ext = Path(file_path).suffix.lower()
    loaders: dict[str, Callable[[], DocumentLoader]] = {
        ".txt": TXTLoader,
        ".pdf": PDFLoader,
        ".docx": DOCXLoader,
    }
    loader_factory = loaders.get(ext)
    if not loader_factory:
        raise ValueError(f"Unsupported file extension: {ext}")
    return loader_factory()
